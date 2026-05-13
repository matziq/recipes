"""Import a recipe from a URL (uses JSON-LD schema.org/Recipe) into OneDrive/Recipes as a .docx.

Usage:
    python import_recipe.py <url> [--category Breads] [--title "Override Title"]
        [--no-image] [--no-mark-new]
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import date
from html import unescape
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

import requests
from docx import Document
from docx.shared import Inches

ONEDRIVE_RECIPES = Path(r"d:/OneDrive/Recipes")
SITE_DIR = Path(__file__).parent
NEW_RECIPES_JSON = SITE_DIR / "new_recipes.json"
UA = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
BROWSER_HEADERS = {
    "User-Agent": UA,
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
    "Accept-Encoding": "gzip, deflate, br",
    "Cache-Control": "no-cache",
    "Pragma": "no-cache",
    "Sec-Ch-Ua": '"Chromium";v="120", "Not_A Brand";v="24"',
    "Sec-Ch-Ua-Mobile": "?0",
    "Sec-Ch-Ua-Platform": '"Windows"',
    "Sec-Fetch-Dest": "document",
    "Sec-Fetch-Mode": "navigate",
    "Sec-Fetch-Site": "none",
    "Sec-Fetch-User": "?1",
    "Upgrade-Insecure-Requests": "1",
}


def strip_html(s: str) -> str:
    s = re.sub(r"<[^>]+>", "", s or "")
    return unescape(s).strip()


def find_recipe(data):
    """Recursively search JSON-LD for a Recipe object."""
    if isinstance(data, dict):
        t = data.get("@type")
        if t == "Recipe" or (isinstance(t, list) and "Recipe" in t):
            return data
        for v in data.values():
            r = find_recipe(v)
            if r:
                return r
    elif isinstance(data, list):
        for v in data:
            r = find_recipe(v)
            if r:
                return r
    return None


def _fetch_html(url: str) -> str:
    sess = requests.Session()
    sess.headers.update(BROWSER_HEADERS)
    parsed = urlparse(url)
    referer = None
    if parsed.scheme and parsed.netloc:
        referer = f"{parsed.scheme}://{parsed.netloc}/"
        sess.headers["Referer"] = referer
    try:
        resp = sess.get(url, timeout=30, allow_redirects=True)
        resp.raise_for_status()
        if "application/ld+json" in resp.text:
            return resp.text
        # Some sites (e.g., onceuponachef.com) serve a JSON-LD-less variant
        # to plain requests clients; treat as a soft failure and fall through.
        soft_failure: Exception | None = None
    except (requests.HTTPError, requests.exceptions.SSLError) as exc:
        soft_failure = exc
        status = None
        if isinstance(exc, requests.HTTPError) and exc.response is not None:
            status = exc.response.status_code
            if status not in (403, 429, 503):
                raise
    # Fallback: curl_cffi impersonates a real Chrome TLS fingerprint to
    # defeat Cloudflare bot-protection challenges and content variants.
    try:
        from curl_cffi import requests as cffi_requests  # type: ignore
    except ImportError as ie:
        raise SystemExit(
            f"Could not fetch {url} ({soft_failure!r}); install curl_cffi to bypass."
        ) from ie
    cffi_headers = {"Accept-Language": "en-US,en;q=0.9"}
    if referer:
        cffi_headers["Referer"] = referer
    last_text = None
    for profile in ("chrome124", "chrome120", "chrome119", "chrome116"):
        try:
            resp2 = cffi_requests.get(
                url, impersonate=profile, timeout=45,
                allow_redirects=True, headers=cffi_headers,
            )
        except Exception:
            continue
        last_text = resp2.text
        if resp2.status_code == 200 and "application/ld+json" in resp2.text:
            return resp2.text
    if last_text and "application/ld+json" in last_text:
        return last_text
    raise SystemExit(
        f"Could not fetch {url}; Cloudflare challenge or JSON-LD content unavailable."
    )


def fetch_recipe(url: str) -> dict:
    html = _fetch_html(url)
    for m in re.finditer(
        r'<script[^>]*type=[\'"]application/ld\+json[\'"][^>]*>(.*?)</script>',
        html, re.DOTALL | re.IGNORECASE,
    ):
        try:
            data = json.loads(m.group(1))
        except json.JSONDecodeError:
            continue
        recipe = find_recipe(data)
        if recipe:
            return recipe
    raise SystemExit("No JSON-LD Recipe schema found on the page.")


def safe_filename(title: str) -> str:
    return re.sub(r'[<>:"/\\|?*]', "", title).strip()


def extract_image_url(image_data) -> str | None:
    """Extract first usable image URL from a schema.org Recipe `image` field."""
    if not image_data:
        return None
    if isinstance(image_data, str):
        return image_data
    if isinstance(image_data, dict):
        url = image_data.get("url") or image_data.get("contentUrl") or image_data.get("@id")
        if isinstance(url, str):
            return url
        return extract_image_url(image_data.get("image"))
    if isinstance(image_data, list):
        for item in image_data:
            url = extract_image_url(item)
            if url:
                return url
    return None


def download_image(url: str) -> BytesIO | None:
    try:
        resp = requests.get(url, headers=BROWSER_HEADERS, timeout=30)
        resp.raise_for_status()
        return BytesIO(resp.content)
    except Exception as exc:
        # Try the same curl_cffi fallback used for HTML fetches.
        try:
            from curl_cffi import requests as cffi_requests  # type: ignore
            r2 = cffi_requests.get(url, impersonate="chrome124", timeout=30)
            if r2.status_code == 200 and r2.content:
                return BytesIO(r2.content)
        except Exception:
            pass
        print(f"Warning: could not download image {url}: {exc}")
        return None


def mark_recipe_as_new(category: str, title: str) -> None:
    """Record this recipe in new_recipes.json so build.py renders a NEW badge."""
    entries: list[dict] = []
    if NEW_RECIPES_JSON.exists():
        try:
            entries = json.loads(NEW_RECIPES_JSON.read_text(encoding="utf-8"))
            if not isinstance(entries, list):
                entries = []
        except json.JSONDecodeError:
            entries = []
    key = (category.strip().lower(), title.strip().lower())
    if any(
        (e.get("category", "").strip().lower(), e.get("title", "").strip().lower()) == key
        for e in entries
    ):
        return
    entries.append({
        "category": category,
        "title": title,
        "added": date.today().isoformat(),
    })
    NEW_RECIPES_JSON.write_text(
        json.dumps(entries, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )


def parse_iso_duration(s: str) -> str:
    """Convert PT1H45M -> '1 hr 45 min'."""
    if not s:
        return ""
    m = re.match(r"PT(?:(\d+)H)?(?:(\d+)M)?", s)
    if not m:
        return s
    h, mn = m.groups()
    parts = []
    if h:
        parts.append(f"{h} hr")
    if mn:
        parts.append(f"{mn} min")
    return " ".join(parts)


def build_docx(recipe: dict, source_url: str, out_path: Path, include_image: bool = True) -> None:
    doc = Document()
    title = strip_html(recipe.get("name", "Untitled Recipe"))
    doc.add_heading(title, level=1)

    # Hero image (optional)
    if include_image:
        img_url = extract_image_url(recipe.get("image"))
        if img_url:
            img_data = download_image(img_url)
            if img_data is not None:
                try:
                    doc.add_picture(img_data, width=Inches(5.0))
                except Exception as exc:
                    print(f"Warning: could not embed image: {exc}")

    desc = strip_html(recipe.get("description", ""))
    if desc:
        doc.add_paragraph(desc)

    # Meta
    meta_bits = []
    yield_v = recipe.get("recipeYield")
    if isinstance(yield_v, list):
        yield_v = yield_v[-1] if yield_v else ""
    if yield_v:
        meta_bits.append(f"Yield: {yield_v}")
    prep = parse_iso_duration(recipe.get("prepTime", ""))
    cook = parse_iso_duration(recipe.get("cookTime", ""))
    total = parse_iso_duration(recipe.get("totalTime", ""))
    if prep:
        meta_bits.append(f"Prep: {prep}")
    if cook:
        meta_bits.append(f"Cook: {cook}")
    if total:
        meta_bits.append(f"Total: {total}")
    if meta_bits:
        p = doc.add_paragraph()
        p.add_run(" • ".join(meta_bits)).italic = True

    # Ingredients
    ings = recipe.get("recipeIngredient") or []
    if ings:
        doc.add_heading("Ingredients", level=2)
        for ing in ings:
            doc.add_paragraph(strip_html(str(ing)), style="List Bullet")

    # Instructions
    instrs = recipe.get("recipeInstructions") or []
    if instrs:
        doc.add_heading("Instructions", level=2)
        steps: list[str] = []
        for it in instrs:
            if isinstance(it, str):
                steps.append(strip_html(it))
            elif isinstance(it, dict):
                if it.get("@type") == "HowToSection":
                    sect = strip_html(it.get("name", ""))
                    if sect:
                        steps.append(f"[{sect}]")
                    for sub in it.get("itemListElement", []) or []:
                        if isinstance(sub, dict):
                            steps.append(strip_html(sub.get("text", "")))
                else:
                    steps.append(strip_html(it.get("text", "")))
        for step in steps:
            if step:
                doc.add_paragraph(step, style="List Number")

    # Source
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"Source: {source_url}")
    r.italic = True

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("url")
    ap.add_argument("--category", default="Breads")
    ap.add_argument("--title", default=None, help="Override title for filename")
    ap.add_argument("--no-image", action="store_true", help="Skip embedding the recipe's main image")
    ap.add_argument("--no-mark-new", action="store_true", help="Do not flag this recipe as NEW for the site")
    args = ap.parse_args()

    recipe = fetch_recipe(args.url)
    title = args.title or strip_html(recipe.get("name", "Untitled"))
    out = ONEDRIVE_RECIPES / args.category / f"{safe_filename(title)}.docx"
    if out.exists():
        print(f"Already exists: {out}")
        if not args.no_mark_new:
            mark_recipe_as_new(args.category, title)
            print(f"Re-marked as NEW: {args.category} / {title}")
        sys.exit(0)
    build_docx(recipe, args.url, out, include_image=not args.no_image)
    print(f"Saved: {out}")
    if not args.no_mark_new:
        mark_recipe_as_new(args.category, title)
        print(f"Marked as NEW: {args.category} / {title}")


if __name__ == "__main__":
    main()
